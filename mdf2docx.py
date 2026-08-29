#!/usr/bin/env python3
"""
mdf2docx - convert University of Sheffield MDF module development form PDFs
into editable Word documents that keep the original layout.

The converter reads the PDF's own geometry (cell borders, fill colours, font
sizes) rather than guessing from the text, so it adapts to whatever each module
form actually contains: different numbers of learning outcomes, assessments,
availability rows and so on.

Usage
-----
  Single file:
      python3 mdf2docx.py form.pdf
      python3 mdf2docx.py form.pdf -o out/form.docx

  Whole folder:
      python3 mdf2docx.py --batch pdfs/ --out docx/

Options
-------
  --body-size PT      body text size in points (default 10)
  --page-breaks       keep the original page breaks (off by default, so
                      sections flow and every heading is spaced identically)
  --font NAME         output font (default Arial)
  --jobs N            parallel workers for batch mode (default 4)

Requires: pdfplumber, python-docx
"""

import argparse
import concurrent.futures
import glob
import os
import re
import sys
import traceback

import pdfplumber
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

# --------------------------------------------------------------------------
# appearance
# --------------------------------------------------------------------------

FONT = "Arial"
INK = "404143"          # body / heading colour used throughout the form
FILL = "ADD8E6"         # pale blue of label and header cells
BLUE_RGB = (0.678431372, 0.847058823, 0.901960784)

PAGE_W = 16838          # A4 landscape, twips
PAGE_H = 11906
M_LEFT, M_RIGHT, M_TOP, M_BOTTOM = 800, 680, 600, 600
CONTENT_W = PAGE_W - M_LEFT - M_RIGHT           # 15358

# cell padding, twips
PAD_NORMAL = dict(top=90, bottom=90, left=110, right=110)
PAD_TIGHT = dict(top=80, bottom=80, left=60, right=60)
TIGHT_FROM_COLS = 6     # tables with this many columns use the tighter padding

# The source PDF is a shrink-to-fit print, so its text is about 6.6pt. Map its
# four sizes onto readable Word sizes, preserving the relative proportions.
SIZE_MAP = [(6.62, 10.0), (8.15, 12.0), (10.19, 15.5), (13.76, 20.0)]
SIZE_TOL = 0.35

# paragraph spacing, points
SP_TITLE_AFTER = 14
SP_H1_BEFORE, SP_H1_AFTER = 20, 11
SP_H2_BEFORE, SP_H2_AFTER = 15, 9
SP_BODY_AFTER = 6
SP_BOX_GAP = 10.3       # gap between the stacked label/value boxes

PDFPLUMBER_TABLE_SETTINGS = {
    "vertical_strategy": "lines",
    "horizontal_strategy": "lines",
    "intersection_tolerance": 2,
    "join_tolerance": 2,
    "snap_tolerance": 2,
}

# --------------------------------------------------------------------------
# Arial (Helvetica-metric) advance widths, units per 1000 em
# --------------------------------------------------------------------------

_PUNCT_R = {' ':278,'!':278,'"':355,'#':556,'$':556,'%':889,'&':667,"'":191,'(':333,
            ')':333,'*':389,'+':584,',':278,'-':333,'.':278,'/':278,':':278,';':278,
            '<':584,'=':584,'>':584,'?':556,'@':1015,'[':278,'\\':278,']':278,'^':469,
            '_':556,'`':333,'{':334,'|':260,'}':334,'~':584}
_PUNCT_B = {' ':278,'!':333,'"':474,'#':556,'$':556,'%':889,'&':722,"'":238,'(':333,
            ')':333,'*':389,'+':584,',':278,'-':333,'.':278,'/':278,':':333,';':333,
            '<':584,'=':584,'>':584,'?':611,'@':975,'[':333,'\\':278,']':333,'^':584,
            '_':556,'`':333,'{':389,'|':280,'}':389,'~':584}
_UPPER_R = [667,667,722,722,667,611,778,722,278,500,667,556,833,722,778,667,778,722,
            667,611,722,667,944,667,667,611]
_UPPER_B = [722,722,722,722,667,611,778,722,278,556,722,611,833,722,778,667,778,722,
            667,611,722,667,944,667,667,611]
_LOWER_R = [556,556,500,556,556,278,556,556,222,222,500,222,833,556,556,556,556,333,
            500,278,556,500,722,500,500,500]
_LOWER_B = [556,611,556,611,556,333,611,611,278,278,556,278,889,611,611,611,611,389,
            556,333,611,556,778,556,556,500]


def _width_table(bold):
    t = dict(_PUNCT_B if bold else _PUNCT_R)
    for i, w in enumerate(_UPPER_B if bold else _UPPER_R):
        t[chr(65 + i)] = w
    for i, w in enumerate(_LOWER_B if bold else _LOWER_R):
        t[chr(97 + i)] = w
    for d in "0123456789":
        t[d] = 556
    return t


_W_REG = _width_table(False)
_W_BLD = _width_table(True)


def text_width(s, size_pt, bold=False):
    """Width of a string in twips at the given point size."""
    tbl = _W_BLD if bold else _W_REG
    units = sum(tbl.get(ch, 556) for ch in str(s))
    return units / 1000.0 * size_pt * 20


def longest_word(s, size_pt, bold=False):
    parts = str(s).split()
    return max((text_width(w, size_pt, bold) for w in parts), default=0.0)


def fit_columns(orig, cols, pad, body_pt):
    """
    Choose column widths that add up to the content width.

    Start from the proportions measured in the PDF, so the layout matches the
    original wherever it can. Any column too narrow for its longest unbreakable
    word (the larger output text makes this common) is widened, and the space
    is taken first from columns with room to spare, then from the widest
    remaining columns. Nothing is changed that does not need to change.

    orig  - relative widths measured from the PDF
    cols  - per column, a list of (text, bold) covering header and body cells
    """
    n = len(orig)
    extra = pad["left"] + pad["right"] + 30
    mins, wants = [], []
    for entries in cols:
        mn = max((longest_word(t, body_pt, b) for t, b in entries), default=0.0)
        wt = max((text_width(t, body_pt, b) for t, b in entries), default=0.0)
        mins.append(int(mn + extra) + 1)
        wants.append(max(int(wt + extra) + 1, mins[-1]))

    if sum(mins) >= CONTENT_W:                     # cannot satisfy all: scale down
        out = [int(v * CONTENT_W / sum(mins)) for v in mins]
        out[-1] += CONTENT_W - sum(out)
        return out

    total = float(sum(orig)) or 1.0
    base = [orig[i] / total * CONTENT_W for i in range(n)]
    out = list(base)

    deficit = [max(0.0, mins[i] - out[i]) for i in range(n)]
    owed = sum(deficit)
    if owed > 0:
        for i in range(n):
            out[i] += deficit[i]
        # stage 1: take from columns holding more than their content needs
        spare = [max(0.0, base[i] - wants[i]) if deficit[i] == 0 else 0.0
                 for i in range(n)]
        take = min(owed, sum(spare))
        if take > 0:
            for i in range(n):
                out[i] -= take * spare[i] / sum(spare)
            owed -= take
        # stage 2: take what is still needed from whatever room is left
        if owed > 0.5:
            room = [max(0.0, out[i] - mins[i]) if deficit[i] == 0 else 0.0
                    for i in range(n)]
            if sum(room) > 0:
                take = min(owed, sum(room))
                for i in range(n):
                    out[i] -= take * room[i] / sum(room)

    out = [int(round(v)) for v in out]
    out[-1] += CONTENT_W - sum(out)                # absorb rounding
    return out


# --------------------------------------------------------------------------
# extraction
# --------------------------------------------------------------------------

def map_size(pdf_size, scale):
    for src, dst in SIZE_MAP:
        if abs(pdf_size - src) <= SIZE_TOL:
            return dst * scale
    return round(pdf_size * 1.5097 * scale * 2) / 2      # unseen size: scale it


def is_blue(rect):
    c = rect.get("non_stroking_color")
    if not c or not isinstance(c, (list, tuple)) or len(c) != 3:
        return False
    return all(abs(a - b) < 0.06 for a, b in zip(c, BLUE_RGB))


def group_lines(words, tol=1.6):
    """Group words into visual lines, in reading order."""
    lines = []
    for w in sorted(words, key=lambda w: (round(w["top"], 1), w["x0"])):
        if lines and abs(w["top"] - lines[-1][0]["top"]) <= tol:
            lines[-1].append(w)
        else:
            lines.append([w])
    return [sorted(ln, key=lambda w: w["x0"]) for ln in lines]


def lines_to_paragraphs(lines):
    """
    Split lines into paragraphs on the vertical gaps. A gap noticeably larger
    than the normal line pitch means a new paragraph rather than a wrap.
    """
    paras = []
    for ln in lines:
        text = " ".join(w["text"] for w in ln)
        size = max(w.get("size", 8) for w in ln)
        bold = sum(1 for w in ln if "Bold" in w.get("fontname", "")) * 2 >= len(ln)
        top, bottom = min(w["top"] for w in ln), max(w["bottom"] for w in ln)
        if paras:
            prev = paras[-1]
            gap = top - prev["bottom"]
            same = abs(size - prev["size"]) < 0.3 and bold == prev["bold"]
            if same and gap < size * 0.9:
                prev["text"] += " " + text
                prev["bottom"] = bottom
                continue
        paras.append(dict(text=text, size=size, bold=bold, top=top, bottom=bottom))
    return paras


def cell_content(words, bbox, pad=1.0):
    """Words whose centre falls inside a cell."""
    x0, top, x1, bottom = bbox
    return [w for w in words
            if x0 - pad <= (w["x0"] + w["x1"]) / 2 <= x1 + pad
            and top - pad <= (w["top"] + w["bottom"]) / 2 <= bottom + pad]


def build_cell(words, bbox, blue_rects):
    inner = cell_content(words, bbox)
    raw = lines_to_paragraphs(group_lines(inner)) if inner else []

    # keep each paragraph's own weight, and how far it sat below the previous
    # one, so a wrapped label does not pick up the spacing of a real gap
    paras = []
    for i, p in enumerate(raw):
        gap = 0.0 if i == 0 else p["top"] - raw[i - 1]["bottom"]
        paras.append(dict(text=p["text"], bold=p["bold"], gap=gap, size=p["size"]))

    x0, top, x1, bottom = bbox
    cx, cy = (x0 + x1) / 2, (top + bottom) / 2
    filled = any(r["x0"] - 1 <= cx <= r["x1"] + 1 and r["top"] - 1 <= cy <= r["bottom"] + 1
                 for r in blue_rects)

    align = "left"
    if inner:
        left_pad = min(w["x0"] for w in inner) - x0
        right_pad = x1 - max(w["x1"] for w in inner)
        if left_pad > 5 and abs(left_pad - right_pad) < 3.5:
            align = "center"

    return dict(paras=paras, fill=filled, align=align, span=1)


def cell_text(cell):
    return " ".join(p["text"] for p in cell["paras"])


def extract_table(page, table, words, blue_rects):
    edges = sorted({round(c[0], 1) for r in table.rows for c in r.cells if c} |
                   {round(c[2], 1) for r in table.rows for c in r.cells if c})
    # collapse edges that are within a hair of each other
    merged = [edges[0]]
    for e in edges[1:]:
        if e - merged[-1] > 2:
            merged.append(e)
    edges = merged
    ncols = len(edges) - 1
    if ncols < 1:
        return None
    widths_pt = [edges[i + 1] - edges[i] for i in range(ncols)]

    rows = []
    for r in table.rows:
        cells, i = [], 0
        raw = list(r.cells)
        while i < len(raw):
            bbox = raw[i]
            if bbox is None:                       # stray empty slot
                cells.append(dict(paras=[], fill=False, align="left", span=1))
                i += 1
                continue
            span = 1
            while i + span < len(raw) and raw[i + span] is None:
                span += 1                          # horizontally merged cell
            c = build_cell(words, bbox, blue_rects)
            c["span"] = span
            cells.append(c)
            i += span
        if cells:
            rows.append(cells)

    if not rows:
        return None
    header = all(c["fill"] for c in rows[0]) and len(rows[0]) > 1
    return dict(kind="table", rows=rows, widths=widths_pt, ncols=ncols,
                header=header, top=table.bbox[1], bottom=table.bbox[3])


def extract_blocks(path):
    """Read the PDF into an ordered list of heading / paragraph / table blocks."""
    blocks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(extra_attrs=["size", "fontname"],
                                       keep_blank_chars=False)
            blue = [r for r in page.rects if is_blue(r)]
            tables = page.find_tables(PDFPLUMBER_TABLE_SETTINGS)
            boxes = [t.bbox for t in tables]

            def inside_table(w):
                cx, cy = (w["x0"] + w["x1"]) / 2, (w["top"] + w["bottom"]) / 2
                return any(b[0] - 2 <= cx <= b[2] + 2 and b[1] - 2 <= cy <= b[3] + 2
                           for b in boxes)

            loose = [w for w in words if not inside_table(w)]
            items = [dict(kind="text", **p) for p in lines_to_paragraphs(group_lines(loose))]
            for t in tables:
                tb = extract_table(page, t, words, blue)
                if tb:
                    items.append(tb)

            items.sort(key=lambda b: b["top"])
            for it in items:
                it["page"] = page.page_number
                blocks.append(it)

    return merge_split_tables(blocks)


def assign_widths(blocks, body_pt):
    """
    Work out column widths once per distinct grid.

    Every stacked label/value box in the form shares one grid in the PDF, so
    they must share one set of fitted widths here too, otherwise each box would
    be sized to its own text and the labels would no longer line up down the
    page. Tables are keyed by their column count and source proportions.
    """
    groups = {}
    for b in blocks:
        if b["kind"] != "table":
            continue
        key = (b["ncols"], tuple(round(w) for w in b["widths"]))
        groups.setdefault(key, []).append(b)

    for (ncols, _), members in groups.items():
        pad = PAD_TIGHT if ncols >= TIGHT_FROM_COLS else PAD_NORMAL
        per_col = [[] for _ in range(ncols)]
        for b in members:
            for row in b["rows"]:
                idx = 0
                for c in row:
                    if c["span"] == 1 and idx < ncols:
                        for p in c["paras"]:
                            per_col[idx].append((p["text"], p["bold"]))
                    idx += c["span"]
        widths = fit_columns(members[0]["widths"], per_col, pad, body_pt)
        for b in members:
            b["fitted"] = widths


def merge_split_tables(blocks):
    """
    Rejoin a data table that the PDF split across a page boundary.

    Only true data tables are merged: the earlier part must have a full header
    row, and the later part must either carry no header or repeat the same one.
    The stacked single-row label boxes never match those conditions, so they
    stay as separate boxes.

    Returns the blocks plus the text of any repeated header rows that were
    deliberately dropped, so the completeness check does not flag them.
    """
    out, dropped = [], []
    for b in blocks:
        if (out and b["kind"] == "table" and out[-1]["kind"] == "table"
                and b["page"] != out[-1]["page"]
                and b["ncols"] == out[-1]["ncols"] and out[-1]["header"]):
            prev = out[-1]
            same_grid = all(abs(a - c) < 3 for a, c in zip(prev["widths"], b["widths"]))
            if same_grid:
                rows = b["rows"]
                if b["header"]:
                    a = [cell_text(c) for c in prev["rows"][0]]
                    c2 = [cell_text(c) for c in rows[0]]
                    if a == c2:
                        dropped.extend(a)           # repeated header, drop it
                        rows = rows[1:]
                    else:
                        out.append(b)
                        continue
                prev["rows"].extend(rows)
                continue
        out.append(b)
    return out, dropped


# --------------------------------------------------------------------------
# docx generation
# --------------------------------------------------------------------------

def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = _el("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        borders.append(_el("w:" + side, val="single", sz="4", space="0", color="000000"))
    tcPr.append(borders)


def set_cell_fill(cell, colour):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(_el("w:shd", val="clear", color="auto", fill=colour))


def set_table_margins(table, pad):
    tblPr = table._tbl.tblPr
    mar = _el("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        mar.append(_el("w:" + side, w=pad[side], type="dxa"))
    tblPr.append(mar)


def set_fixed_layout(table, widths):
    table.autofit = False
    table._tbl.tblPr.append(_el("w:tblLayout", type="fixed"))
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        grid.append(_el("w:gridCol", w=w))
    table._tbl.insert(1, grid)


def style_run(run, size_pt, bold, font):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(INK)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "cs"):
        rfonts.set(qn("w:" + attr), font)


def add_paragraph(container, text, size_pt, bold=False, align="left",
                  before=0, after=0, keep=False, font=FONT):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = keep
    if align == "center":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        style_run(p.add_run(text), size_pt, bold, font)
    return p


def add_gap(doc, font=FONT):
    """Thin spacer that also stops two adjacent tables merging in Word."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(SP_BOX_GAP)
    style_run(p.add_run(""), 5, False, font)
    return p


def render_table(doc, block, body_pt, font):
    ncols = block["ncols"]
    pad = PAD_TIGHT if ncols >= TIGHT_FROM_COLS else PAD_NORMAL
    widths = block["fitted"]

    table = doc.add_table(rows=0, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_layout(table, widths)
    set_table_margins(table, pad)

    for r_i, row in enumerate(block["rows"]):
        docx_row = table.add_row()
        # a boxed row that breaks across a page leaves a stray empty half behind
        docx_row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        if r_i == 0 and block["header"] and len(block["rows"]) > 2:
            docx_row._tr.get_or_add_trPr().append(_el("w:tblHeader", val="true"))
        cells = docx_row.cells
        idx = 0
        for c in row:
            cell = cells[idx]
            if c["span"] > 1:
                cell = cell.merge(cells[min(idx + c["span"] - 1, ncols - 1)])
            cell.width = Twips(sum(widths[idx:idx + c["span"]]))
            cell.vertical_alignment = (WD_ALIGN_VERTICAL.CENTER
                                       if (r_i == 0 and block["header"])
                                       else WD_ALIGN_VERTICAL.TOP)
            set_cell_borders(cell)
            if c["fill"]:
                set_cell_fill(cell, FILL)

            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            paras = c["paras"] or [dict(text="", bold=False, gap=0.0, size=body_pt)]
            for p_i, p in enumerate(paras):
                nxt = paras[p_i + 1] if p_i + 1 < len(paras) else None
                # only carry spacing where the PDF had a real gap, not a wrap
                after = 0
                if nxt is not None and nxt["gap"] > p["size"] * 0.55:
                    after = SP_BODY_AFTER * (body_pt / 10.0)
                add_paragraph(cell, p["text"], body_pt, bold=p["bold"],
                              align=c["align"], after=after, font=font)
            idx += c["span"]
    return table


def build_document(blocks, body_pt=10.0, font=FONT, page_breaks=False):
    scale = body_pt / 10.0
    assign_widths(blocks, body_pt)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(body_pt)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(0)

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Twips(PAGE_W), Twips(PAGE_H)
    sec.left_margin, sec.right_margin = Twips(M_LEFT), Twips(M_RIGHT)
    sec.top_margin, sec.bottom_margin = Twips(M_TOP), Twips(M_BOTTOM)

    prev_page = blocks[0]["page"] if blocks else 1
    for i, b in enumerate(blocks):
        if page_breaks and b["page"] != prev_page:
            doc.add_page_break()
        prev_page = b["page"]

        if b["kind"] == "table":
            render_table(doc, b, body_pt, font)
            nxt = blocks[i + 1] if i + 1 < len(blocks) else None
            if nxt is not None and nxt["kind"] == "table":
                add_gap(doc, font)                 # keeps the boxes separate
            continue

        size = map_size(b["size"], scale)
        if b["size"] >= 13.0:                      # document title
            add_paragraph(doc, b["text"], size, bold=True,
                          after=SP_TITLE_AFTER * scale, font=font)
        elif b["size"] >= 9.5:                     # section heading
            add_paragraph(doc, b["text"], size, bold=True,
                          before=(0 if i == 0 else SP_H1_BEFORE * scale),
                          after=SP_H1_AFTER * scale, keep=True, font=font)
        elif b["size"] >= 7.5:                     # sub heading
            add_paragraph(doc, b["text"], size, bold=True,
                          before=SP_H2_BEFORE * scale, after=SP_H2_AFTER * scale,
                          keep=True, font=font)
        else:                                      # body text
            add_paragraph(doc, b["text"], size, bold=b["bold"],
                          after=SP_BODY_AFTER * scale, font=font)
    return doc


# --------------------------------------------------------------------------
# driver
# --------------------------------------------------------------------------

def _norm(text):
    return re.sub(r"[\u2018\u2019]", "'", text or "")


def _tokens(text):
    return [t for t in re.split(r"\s+", _norm(text)) if t]


def docx_tokens(path):
    doc = Document(path)
    out = []
    for p in doc.paragraphs:
        out += _tokens(p.text)
    for t in doc.tables:
        for row in t.rows:
            seen = set()
            for c in row.cells:
                if c._tc in seen:                  # merged cells repeat
                    continue
                seen.add(c._tc)
                out += _tokens(c.text)
    return out


def verify(pdf_path, docx_path, dropped=()):
    """
    Compare the words in the source PDF with the words in the output.

    Cheap insurance on a large batch: if a table or cell is ever missed, this
    reports it rather than letting a quietly incomplete document through.
    """
    from collections import Counter
    src = Counter()
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            for w in page.extract_words():
                src[_norm(w["text"])] += 1
    out = Counter(docx_tokens(docx_path))
    for text in dropped:                           # repeated headers we removed
        for tok in _tokens(text):
            out[tok] += 1
    missing = src - out
    total = sum(src.values()) or 1
    lost = sum(missing.values())
    return dict(total=total, missing=lost,
                coverage=100.0 * (total - lost) / total,
                sample=[w for w, _ in missing.most_common(8)])


def convert(pdf_path, out_path, body_pt=10.0, font=FONT, page_breaks=False,
            check=True):
    blocks, dropped = extract_blocks(pdf_path)
    if not blocks:
        raise ValueError("no content found - is this an MDF form?")
    doc = build_document(blocks, body_pt=body_pt, font=font, page_breaks=page_breaks)
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    n_tables = sum(1 for b in blocks if b["kind"] == "table")
    result = dict(source=pdf_path, output=out_path, blocks=len(blocks), tables=n_tables)
    result["check"] = verify(pdf_path, out_path, dropped) if check else None
    return result


def _job(args):
    src, dest, body_pt, font, page_breaks, check = args
    try:
        r = convert(src, dest, body_pt, font, page_breaks, check)
        return True, src, r
    except Exception as exc:                       # keep the batch running
        return False, src, "%s: %s" % (type(exc).__name__, exc)


def main(argv=None):
    ap = argparse.ArgumentParser(description="Convert MDF form PDFs to Word documents.")
    ap.add_argument("input", nargs="?", help="a PDF file, or a folder with --batch")
    ap.add_argument("--batch", metavar="DIR", help="convert every PDF in this folder")
    ap.add_argument("-o", "--out", help="output file, or output folder in batch mode")
    ap.add_argument("--body-size", type=float, default=10.0,
                    help="body text size in points (default 10)")
    ap.add_argument("--font", default=FONT, help="output font (default Arial)")
    ap.add_argument("--page-breaks", action="store_true",
                    help="keep the PDF's original page breaks")
    ap.add_argument("--jobs", type=int, default=4, help="parallel workers in batch mode")
    ap.add_argument("--recursive", action="store_true", help="search subfolders too")
    ap.add_argument("--no-check", action="store_true",
                    help="skip the check that all PDF text reached the document")
    args = ap.parse_args(argv)

    if args.batch:
        pattern = "**/*.pdf" if args.recursive else "*.pdf"
        sources = sorted(glob.glob(os.path.join(args.batch, pattern),
                                   recursive=args.recursive))
        if not sources:
            print("No PDFs found in %s" % args.batch, file=sys.stderr)
            return 1
        out_dir = args.out or os.path.join(args.batch, "docx")
        jobs = []
        for src in sources:
            rel = os.path.relpath(src, args.batch)
            dest = os.path.join(out_dir, os.path.splitext(rel)[0] + ".docx")
            jobs.append((src, dest, args.body_size, args.font, args.page_breaks,
                         not args.no_check))

        ok, failed, warned = 0, [], []
        workers = max(1, min(args.jobs, len(jobs)))
        with concurrent.futures.ProcessPoolExecutor(max_workers=workers) as pool:
            for i, (good, src, info) in enumerate(pool.map(_job, jobs), 1):
                name = os.path.basename(src)
                if good:
                    ok += 1
                    chk = info.get("check")
                    note = ""
                    if chk and chk["missing"]:
                        note = "  WARNING %.1f%% of text matched, missing e.g. %s" % (
                            chk["coverage"], ", ".join(chk["sample"][:4]))
                        warned.append((name, note.strip()))
                    print("[%d/%d] %s  (%d tables)%s" % (i, len(jobs), name,
                                                         info["tables"], note))
                else:
                    failed.append((name, info))
                    print("[%d/%d] %s  FAILED - %s" % (i, len(jobs), name, info))

        print("\nConverted %d of %d. Output in %s" % (ok, len(jobs), out_dir))
        if warned:
            print("Check these (some text did not carry over):")
            for name, why in warned:
                print("  %s  %s" % (name, why))
        if failed:
            print("Failed:")
            for name, why in failed:
                print("  %s  %s" % (name, why))
        return 0 if not failed else 2

    if not args.input:
        ap.error("give a PDF file, or use --batch FOLDER")
    dest = args.out or os.path.splitext(args.input)[0] + ".docx"
    if os.path.isdir(dest):
        dest = os.path.join(dest, os.path.splitext(os.path.basename(args.input))[0] + ".docx")
    try:
        r = convert(args.input, dest, args.body_size, args.font, args.page_breaks,
                    not args.no_check)
    except Exception:
        traceback.print_exc()
        return 1
    print("Wrote %s  (%d tables)" % (r["output"], r["tables"]))
    chk = r.get("check")
    if chk:
        if chk["missing"]:
            print("WARNING: %.1f%% of the PDF text matched. Missing e.g. %s"
                  % (chk["coverage"], ", ".join(chk["sample"])))
        else:
            print("All %d words from the PDF are present." % chk["total"])
    return 0


if __name__ == "__main__":
    sys.exit(main())
